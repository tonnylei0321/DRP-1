#!/usr/bin/env python3
"""
Design Document Generator - 将 Markdown 设计文档转换为 Word (.docx) 和 XML 格式

基于 design-template 目录下的模板，将 Markdown 设计文档转换为企业标准格式的文档。

使用方法:
    python design_doc_generator.py <markdown_file> [options]

Options:
    --output-dir <dir>        输出目录（默认创建以文档名命名的子目录）
    --template <docx>         Word 模板文件（可选）
    --author <name>           文档作者
    --version <ver>           文档版本号
    --doc-number <num>        文档编号
    --format [docx|xml|both]  输出格式（默认: both）
"""

import argparse
import json
import os
import re
import sys
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from xml.dom import minidom


@dataclass
class DocumentSection:
    """文档章节"""
    level: int
    title: str
    content: List[str] = field(default_factory=list)
    subsections: List['DocumentSection'] = field(default_factory=list)
    tables: List[List[List[str]]] = field(default_factory=list)
    code_blocks: List[Tuple[str, str]] = field(default_factory=list)  # (language, code)


@dataclass 
class DesignDocument:
    """设计文档结构"""
    title: str
    description: str = ""
    author: str = ""
    version: str = "1.0"
    doc_number: str = ""
    created_date: str = ""
    modified_date: str = ""
    sections: List[DocumentSection] = field(default_factory=list)
    


class MarkdownParser:
    """Markdown 文档解析器"""
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        with open(self.file_path, 'r', encoding='utf-8') as f:
            self.content = f.read()
        
        self.lines = self.content.split('\n')
    
    def parse(self) -> DesignDocument:
        """解析 Markdown 文档"""
        doc = DesignDocument(
            title=self._extract_title(),
            description=self._extract_description(),
            created_date=datetime.now().strftime("%Y-%m-%d"),
            modified_date=datetime.now().strftime("%Y-%m-%d")
        )
        
        doc.sections = self._parse_sections()
        return doc
    
    def _extract_title(self) -> str:
        """提取文档标题（第一个 # 标题）"""
        for line in self.lines:
            if line.startswith('# ') and not line.startswith('## '):
                return line[2:].strip()
        return self.file_path.stem
    
    def _extract_description(self) -> str:
        """提取文档描述（标题下的 > 引用块）"""
        in_description = False
        description_lines = []
        
        for line in self.lines:
            if line.startswith('# ') and not line.startswith('## '):
                in_description = True
                continue
            if in_description:
                if line.startswith('>'):
                    description_lines.append(line[1:].strip())
                elif line.strip() == '':
                    continue
                elif line.startswith('#'):
                    break
                else:
                    break
        
        return ' '.join(description_lines)
    
    def _parse_sections(self) -> List[DocumentSection]:
        """解析所有章节"""
        sections = []
        current_section = None
        current_subsection = None
        current_subsubsection = None
        
        i = 0
        while i < len(self.lines):
            line = self.lines[i]
            
            # 检测标题级别
            if line.startswith('## ') and not line.startswith('### '):
                # 二级标题 - 新章节
                if current_section:
                    sections.append(current_section)
                current_section = DocumentSection(
                    level=2,
                    title=self._clean_title(line[3:])
                )
                current_subsection = None
                current_subsubsection = None
            
            elif line.startswith('### ') and not line.startswith('#### '):
                # 三级标题 - 子章节
                if current_section:
                    current_subsection = DocumentSection(
                        level=3,
                        title=self._clean_title(line[4:])
                    )
                    current_section.subsections.append(current_subsection)
                    current_subsubsection = None
            
            elif line.startswith('#### '):
                # 四级标题 - 子子章节
                if current_subsection:
                    current_subsubsection = DocumentSection(
                        level=4,
                        title=self._clean_title(line[5:])
                    )
                    current_subsection.subsections.append(current_subsubsection)
            
            elif line.startswith('```'):
                # 代码块
                lang = line[3:].strip()
                code_lines = []
                i += 1
                while i < len(self.lines) and not self.lines[i].startswith('```'):
                    code_lines.append(self.lines[i])
                    i += 1
                code = '\n'.join(code_lines)
                
                target = current_subsubsection or current_subsection or current_section
                if target:
                    target.code_blocks.append((lang, code))
            
            elif line.startswith('|'):
                # 表格
                table_lines = [line]
                i += 1
                while i < len(self.lines) and self.lines[i].startswith('|'):
                    table_lines.append(self.lines[i])
                    i += 1
                i -= 1  # 回退一行
                
                table = self._parse_table(table_lines)
                target = current_subsubsection or current_subsection or current_section
                if target and table:
                    target.tables.append(table)
            
            else:
                # 普通内容
                if line.strip() and not line.startswith('#') and not line.startswith('---'):
                    target = current_subsubsection or current_subsection or current_section
                    if target:
                        target.content.append(line.strip())
            
            i += 1
        
        # 添加最后一个章节
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _clean_title(self, title: str) -> str:
        """清理标题（去除序号等）"""
        title = title.strip()
        # 去除 "1.2.3 xxx" 格式
        title = re.sub(r'^\d+(\.\d+)*\s+', '', title)
        # 去除 "1. xxx" 格式
        title = re.sub(r'^\d+\.\s*', '', title)
        return title.strip()
    
    def _parse_table(self, lines: List[str]) -> List[List[str]]:
        """解析 Markdown 表格"""
        if len(lines) < 2:
            return []
        
        rows = []
        for i, line in enumerate(lines):
            # 跳过分隔行
            if re.match(r'^\|[\s\-:|]+\|$', line):
                continue
            
            # 解析单元格
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
        
        return rows


class WordDocumentGenerator:
    """Word 文档生成器 - 符合用友微服务详细设计模板 v1.6 规范"""

    def __init__(self, template_path: Optional[str] = None):
        self.template_path = template_path

    def generate(self, doc: DesignDocument, output_path: str) -> str:
        """生成 Word 文档"""
        from docx import Document
        from docx.shared import Inches, Pt, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 创建文档
        word_doc = Document()

        # 设置页面为 A4，匹配模板
        self._setup_page(word_doc)

        # 设置默认样式字体
        self._setup_styles(word_doc)

        # 添加封面
        self._add_cover_page(word_doc, doc)

        # 添加修订记录表
        self._add_revision_table(word_doc, doc)

        # 添加目录占位
        self._add_toc_placeholder(word_doc)

        # 添加内容章节（带自动编号）
        for i, section in enumerate(doc.sections, 1):
            self._add_section(word_doc, section, str(i))

        # 保存文档
        word_doc.save(output_path)
        return output_path

    def _setup_page(self, word_doc):
        """设置页面尺寸和边距（匹配模板 A4）"""
        from docx.shared import Cm, Emu

        section = word_doc.sections[0]
        # A4 尺寸
        section.page_width = Emu(7560310)   # 21cm
        section.page_height = Emu(10692130) # 29.7cm
        # 边距：上下 1 英寸，左右 1.25 英寸
        section.top_margin = Emu(914400)
        section.bottom_margin = Emu(914400)
        section.left_margin = Emu(1143000)
        section.right_margin = Emu(1143000)
        # 页眉页脚距离
        section.header_distance = Emu(540385)
        section.footer_distance = Emu(629920)

    def _setup_styles(self, word_doc):
        """设置文档样式以匹配模板"""
        from docx.shared import Pt, Emu, RGBColor
        from docx.oxml.ns import qn

        # Normal 样式 - Times New Roman
        normal_style = word_doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(10.5)  # 五号
        normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # Heading 1 - 22pt 粗体
        h1 = word_doc.styles['Heading 1']
        h1.font.size = Emu(279400)  # 22pt
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h1.paragraph_format.space_before = Emu(215900)
        h1.paragraph_format.space_after = Emu(209550)

        # Heading 2 - Arial 16pt 粗体
        h2 = word_doc.styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Emu(203200)  # 16pt
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h2.paragraph_format.space_before = Emu(228600)
        h2.paragraph_format.space_after = Emu(228600)

        # Heading 3 - 14pt 粗体
        h3 = word_doc.styles['Heading 3']
        h3.font.size = Emu(177800)  # 14pt
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0, 0, 0)
        h3.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h3.paragraph_format.space_before = Emu(165100)
        h3.paragraph_format.space_after = Emu(165100)

    def _add_cover_page(self, word_doc, doc: DesignDocument):
        """添加封面（匹配模板格式）"""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 空行（模拟模板的封面顶部留白）
        for _ in range(6):
            word_doc.add_paragraph()

        # 主标题 - 36pt 粗体 居中 黑色
        title_para = word_doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(doc.title)
        title_run.bold = True
        title_run.font.size = Pt(36)
        title_run.font.color.rgb = RGBColor(0, 0, 0)

        # 空行
        word_doc.add_paragraph()

        # 副标题 - 26pt 粗体 居中 黑色
        subtitle = word_doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run("详细设计文档")
        sub_run.bold = True
        sub_run.font.size = Pt(26)
        sub_run.font.color.rgb = RGBColor(0, 0, 0)

        # 空行（封面中间留白）
        for _ in range(8):
            word_doc.add_paragraph()

        # 文档信息 - 14pt 粗体 左对齐（匹配模板）
        info_items = [
            ("文档编号", doc.doc_number or ""),
            ("创建日期", doc.created_date),
            ("最后修改日期", doc.modified_date),
            ("版 本 号", doc.version),
            ("电子版文件名", f"{Path(doc.title).stem if doc.title else ''}.docx"),
        ]

        for label, value in info_items:
            para = word_doc.add_paragraph()
            run = para.add_run(f"{label}：{value}")
            run.bold = True
            run.font.size = Pt(14)

        # 分页
        word_doc.add_page_break()

    def _add_revision_table(self, word_doc, doc: DesignDocument):
        """添加修订记录表（匹配模板 5 列格式）"""
        from docx.shared import Pt

        # 标题使用 Normal 样式（模板中不是 Heading）
        title_para = word_doc.add_paragraph()
        title_run = title_para.add_run("文档修订摘要")
        title_run.bold = True
        title_run.font.size = Pt(16)

        # 5 列修订表（匹配模板：日期/修订号/描述/著者/审阅者）
        table = word_doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'

        # 表头
        headers = ["日期", "修订号", "描述", "著者", "审阅者"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # 初始版本记录
        table.rows[1].cells[0].text = doc.created_date
        table.rows[1].cells[1].text = doc.version
        table.rows[1].cells[2].text = "初始版本"
        table.rows[1].cells[3].text = doc.author or "AI Coding Team"
        table.rows[1].cells[4].text = ""

        word_doc.add_page_break()

    def _add_toc_placeholder(self, word_doc):
        """添加目录占位符"""
        from docx.shared import Pt

        # 目录标题使用 Normal 样式（匹配模板）
        title_para = word_doc.add_paragraph()
        title_run = title_para.add_run("目录")
        title_run.bold = True
        title_run.font.size = Pt(16)

        para = word_doc.add_paragraph()
        para.add_run('（目录将在 Word 中自动生成，请使用「引用 > 目录」功能）')
        word_doc.add_page_break()

    def _add_section(self, word_doc, section: DocumentSection, number: str):
        """添加章节（带编号，匹配企业模板目录格式）"""
        # 标题带编号前缀（如 "1 概述"、"1.1 目的"、"1.1.1 实体模型"）
        heading_text = f"{number}\t{section.title}"
        word_doc.add_heading(heading_text, level=section.level - 1)

        # 内容
        for content in section.content:
            # 清理 Markdown 格式
            clean_content = self._clean_markdown(content)
            if clean_content:
                word_doc.add_paragraph(clean_content)

        # 代码块
        for lang, code in section.code_blocks:
            from docx.shared import Pt
            para = word_doc.add_paragraph()
            para.style = 'No Spacing'
            run = para.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        # 表格
        for table_data in section.tables:
            if table_data:
                self._add_table(word_doc, table_data)

        # 子章节
        for i, subsection in enumerate(section.subsections, 1):
            self._add_section(word_doc, subsection, f"{number}.{i}")

    def _add_table(self, word_doc, table_data: List[List[str]]):
        """添加表格"""
        if not table_data:
            return

        rows = len(table_data)
        cols = max(len(row) for row in table_data)

        table = word_doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                if j < cols:
                    table.rows[i].cells[j].text = self._clean_markdown(cell)
                    # 表头加粗
                    if i == 0:
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        # 表格后空行
        word_doc.add_paragraph()

    def _clean_markdown(self, text: str) -> str:
        """清理 Markdown 格式"""
        # 去除粗体
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # 去除斜体
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        # 去除代码标记
        text = re.sub(r'`(.*?)`', r'\1', text)
        # 去除链接，保留文本
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        # 去除列表标记
        text = re.sub(r'^[\-\*]\s+', '', text)
        return text.strip()


class XMLDocumentGenerator:
    """XML 文档生成器"""
    
    def generate(self, doc: DesignDocument, output_path: str) -> str:
        """生成 XML 文档"""
        root = ET.Element("DesignDocument")
        root.set("version", "1.0")
        root.set("encoding", "UTF-8")
        
        # 元数据
        metadata = ET.SubElement(root, "Metadata")
        ET.SubElement(metadata, "Title").text = doc.title
        ET.SubElement(metadata, "Description").text = doc.description
        ET.SubElement(metadata, "Author").text = doc.author or "AI Coding Team"
        ET.SubElement(metadata, "Version").text = doc.version
        ET.SubElement(metadata, "DocNumber").text = doc.doc_number or ""
        ET.SubElement(metadata, "CreatedDate").text = doc.created_date
        ET.SubElement(metadata, "ModifiedDate").text = doc.modified_date
        
        # 章节
        sections_elem = ET.SubElement(root, "Sections")
        for i, section in enumerate(doc.sections, 1):
            self._add_section_xml(sections_elem, section, str(i))
        
        # 格式化输出
        xml_str = minidom.parseString(ET.tostring(root, encoding='unicode')).toprettyxml(indent="  ")
        
        # 移除空行
        xml_lines = [line for line in xml_str.split('\n') if line.strip()]
        xml_str = '\n'.join(xml_lines)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(xml_str)
        
        return output_path
    
    def _add_section_xml(self, parent: ET.Element, section: DocumentSection, number: str):
        """添加章节 XML"""
        section_elem = ET.SubElement(parent, "Section")
        section_elem.set("number", number)
        section_elem.set("level", str(section.level))
        
        ET.SubElement(section_elem, "Title").text = section.title
        
        # 内容
        if section.content:
            content_elem = ET.SubElement(section_elem, "Content")
            for para in section.content:
                ET.SubElement(content_elem, "Paragraph").text = para
        
        # 代码块
        if section.code_blocks:
            code_elem = ET.SubElement(section_elem, "CodeBlocks")
            for lang, code in section.code_blocks:
                block = ET.SubElement(code_elem, "CodeBlock")
                block.set("language", lang or "text")
                block.text = code
        
        # 表格
        if section.tables:
            tables_elem = ET.SubElement(section_elem, "Tables")
            for table_data in section.tables:
                table_elem = ET.SubElement(tables_elem, "Table")
                for row in table_data:
                    row_elem = ET.SubElement(table_elem, "Row")
                    for cell in row:
                        ET.SubElement(row_elem, "Cell").text = cell
        
        # 子章节
        if section.subsections:
            subsections_elem = ET.SubElement(section_elem, "Subsections")
            for i, subsection in enumerate(section.subsections, 1):
                self._add_section_xml(subsections_elem, subsection, f"{number}.{i}")


def generate_design_document(
    markdown_file: str,
    output_dir: Optional[str] = None,
    template_path: Optional[str] = None,
    author: str = "",
    version: str = "1.0",
    doc_number: str = "",
    output_format: str = "both"
) -> Dict[str, str]:
    """
    生成设计文档
    
    Args:
        markdown_file: Markdown 源文件路径
        output_dir: 输出目录（默认创建以文档名命名的子目录）
        template_path: Word 模板文件路径
        author: 文档作者
        version: 文档版本
        doc_number: 文档编号
        output_format: 输出格式 (docx, xml, both)
    
    Returns:
        生成的文件路径字典
    """
    # 解析 Markdown
    parser = MarkdownParser(markdown_file)
    doc = parser.parse()
    
    # 设置元数据
    doc.author = author
    doc.version = version
    doc.doc_number = doc_number
    
    # 确定输出目录
    md_path = Path(markdown_file)
    if output_dir:
        out_dir = Path(output_dir)
    else:
        # 默认在 docs/archive/<文档名>-<日期> 下创建
        timestamp = datetime.now().strftime("%Y%m%d")
        out_dir = md_path.parent / "archive" / f"{md_path.stem}-{timestamp}"
    
    out_dir.mkdir(parents=True, exist_ok=True)
    
    results = {"output_dir": str(out_dir)}
    
    # 生成 Word 文档
    if output_format in ("docx", "both"):
        docx_path = out_dir / f"{md_path.stem}.docx"
        word_gen = WordDocumentGenerator(template_path)
        word_gen.generate(doc, str(docx_path))
        results["docx"] = str(docx_path)
    
    # 生成 XML 文档
    if output_format in ("xml", "both"):
        xml_path = out_dir / f"{md_path.stem}.xml"
        xml_gen = XMLDocumentGenerator()
        xml_gen.generate(doc, str(xml_path))
        results["xml"] = str(xml_path)
    
    # 复制原始 Markdown
    import shutil
    md_copy = out_dir / md_path.name
    shutil.copy(markdown_file, md_copy)
    results["markdown"] = str(md_copy)
    
    # 生成索引文件
    index_path = out_dir / "index.json"
    index_data = {
        "title": doc.title,
        "description": doc.description,
        "author": doc.author,
        "version": doc.version,
        "doc_number": doc.doc_number,
        "created_date": doc.created_date,
        "files": {
            "markdown": md_path.name,
            "docx": f"{md_path.stem}.docx" if output_format in ("docx", "both") else None,
            "xml": f"{md_path.stem}.xml" if output_format in ("xml", "both") else None,
        },
        "sections": [s.title for s in doc.sections]
    }
    
    with open(index_path, 'w', encoding='utf-8') as f:
        json.dump(index_data, f, ensure_ascii=False, indent=2)
    results["index"] = str(index_path)
    
    return results


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description="将 Markdown 设计文档转换为 Word (.docx) 和 XML 格式",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    # 基本使用
    python design_doc_generator.py docs/aicoding-scaffold-design.md
    
    # 指定输出目录
    python design_doc_generator.py docs/design.md --output-dir ./output
    
    # 使用模板
    python design_doc_generator.py docs/design.md --template design-template/xxx微服务详细设计参考模板-v1.6.docx
    
    # 只生成 Word
    python design_doc_generator.py docs/design.md --format docx
    
    # 指定文档信息
    python design_doc_generator.py docs/design.md --author "张三" --version "2.0" --doc-number "DOC-2026-001"
        """
    )
    
    parser.add_argument("file", help="Markdown 设计文档路径")
    parser.add_argument(
        "--output-dir", "-o",
        help="输出目录（默认创建以文档名命名的子目录）"
    )
    parser.add_argument(
        "--template", "-t",
        help="Word 模板文件路径"
    )
    parser.add_argument(
        "--author", "-a",
        default="",
        help="文档作者"
    )
    parser.add_argument(
        "--version", "-v",
        default="1.0",
        help="文档版本号 (默认: 1.0)"
    )
    parser.add_argument(
        "--doc-number", "-n",
        default="",
        help="文档编号"
    )
    parser.add_argument(
        "--format", "-f",
        choices=["docx", "xml", "both"],
        default="both",
        help="输出格式 (默认: both)"
    )
    
    args = parser.parse_args()
    
    try:
        results = generate_design_document(
            markdown_file=args.file,
            output_dir=args.output_dir,
            template_path=args.template,
            author=args.author,
            version=args.version,
            doc_number=args.doc_number,
            output_format=args.format
        )
        
        print("✓ 文档生成成功！")
        print(f"  输出目录: {results['output_dir']}")
        if 'docx' in results:
            print(f"  Word 文档: {results['docx']}")
        if 'xml' in results:
            print(f"  XML 文档: {results['xml']}")
        print(f"  索引文件: {results['index']}")
    
    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
