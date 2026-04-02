#!/usr/bin/env python3
"""
Document Reader - 读取 Word (.docx) 和 Excel (.xlsx/.xls) 文件内容

支持的格式:
- Word: .docx
- Excel: .xlsx, .xls

使用方法:
    python doc_reader.py <file_path> [options]

Options:
    --format [text|json|markdown]  输出格式 (默认: markdown)
    --sheet <name>                 指定 Excel 工作表名称
    --all-sheets                   读取所有工作表
    --output <file>                输出到文件
"""

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Union


def check_dependencies() -> Dict[str, bool]:
    """检查必要的依赖库是否已安装"""
    deps = {}
    try:
        import docx
        deps['python-docx'] = True
    except ImportError:
        deps['python-docx'] = False
    
    try:
        import openpyxl
        deps['openpyxl'] = True
    except ImportError:
        deps['openpyxl'] = False
    
    try:
        import xlrd
        deps['xlrd'] = True
    except ImportError:
        deps['xlrd'] = False
    
    return deps


def install_hint(missing_deps: List[str]) -> str:
    """返回安装提示"""
    return f"请安装缺失的依赖: pip install {' '.join(missing_deps)}"


class WordReader:
    """Word 文档读取器"""
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        if self.file_path.suffix.lower() != '.docx':
            raise ValueError(f"不支持的文件格式: {self.file_path.suffix}")
    
    def read(self) -> Dict[str, Any]:
        """读取 Word 文档内容"""
        from docx import Document
        from docx.table import Table
        
        doc = Document(str(self.file_path))
        
        result = {
            "file_name": self.file_path.name,
            "file_type": "docx",
            "paragraphs": [],
            "tables": [],
            "headings": [],
            "metadata": {}
        }
        
        # 提取核心属性
        core_props = doc.core_properties
        result["metadata"] = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }
        
        # 提取段落和标题
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检查是否为标题
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '').replace('Heading', '1')
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                result["headings"].append({
                    "level": level,
                    "text": text
                })
            
            result["paragraphs"].append({
                "style": para.style.name,
                "text": text
            })
        
        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "index": table_idx,
                "rows": []
            }
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data["rows"].append(row_data)
            result["tables"].append(table_data)
        
        return result
    
    def to_markdown(self) -> str:
        """将 Word 文档转换为 Markdown 格式"""
        data = self.read()
        lines = []
        
        # 文件标题
        lines.append(f"# {data['file_name']}")
        lines.append("")
        
        # 元数据
        if any(data["metadata"].values()):
            lines.append("## 文档信息")
            for key, value in data["metadata"].items():
                if value:
                    lines.append(f"- **{key}**: {value}")
            lines.append("")
        
        # 内容
        lines.append("## 内容")
        lines.append("")
        
        for para in data["paragraphs"]:
            style = para["style"]
            text = para["text"]
            
            if style.startswith("Heading"):
                level = style.replace("Heading ", "").replace("Heading", "1")
                try:
                    level = int(level) + 1  # 因为文件名已经是 H1
                except ValueError:
                    level = 2
                lines.append(f"{'#' * level} {text}")
            else:
                lines.append(text)
            lines.append("")
        
        # 表格
        if data["tables"]:
            lines.append("## 表格")
            lines.append("")
            
            for table in data["tables"]:
                if table["rows"]:
                    # 表头
                    header = table["rows"][0]
                    lines.append("| " + " | ".join(header) + " |")
                    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                    
                    # 数据行
                    for row in table["rows"][1:]:
                        # 确保行数据与表头列数匹配
                        padded_row = row + [""] * (len(header) - len(row))
                        lines.append("| " + " | ".join(padded_row[:len(header)]) + " |")
                    lines.append("")
        
        return "\n".join(lines)
    
    def to_text(self) -> str:
        """将 Word 文档转换为纯文本"""
        data = self.read()
        lines = []
        
        lines.append(f"文件: {data['file_name']}")
        lines.append("=" * 50)
        lines.append("")
        
        for para in data["paragraphs"]:
            lines.append(para["text"])
            lines.append("")
        
        if data["tables"]:
            lines.append("=" * 50)
            lines.append("表格内容:")
            lines.append("")
            
            for table in data["tables"]:
                for row in table["rows"]:
                    lines.append(" | ".join(row))
                lines.append("-" * 30)
        
        return "\n".join(lines)


class ExcelReader:
    """Excel 文档读取器"""
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = self.file_path.suffix.lower()
        if suffix not in ['.xlsx', '.xls']:
            raise ValueError(f"不支持的文件格式: {suffix}")
        
        self.is_xlsx = suffix == '.xlsx'
    
    def read(self, sheet_name: Optional[str] = None, all_sheets: bool = False) -> Dict[str, Any]:
        """读取 Excel 文档内容"""
        result = {
            "file_name": self.file_path.name,
            "file_type": "xlsx" if self.is_xlsx else "xls",
            "sheets": {}
        }
        
        if self.is_xlsx:
            result["sheets"] = self._read_xlsx(sheet_name, all_sheets)
        else:
            result["sheets"] = self._read_xls(sheet_name, all_sheets)
        
        return result
    
    def _read_xlsx(self, sheet_name: Optional[str], all_sheets: bool) -> Dict[str, List[List[Any]]]:
        """读取 .xlsx 文件"""
        from openpyxl import load_workbook
        
        wb = load_workbook(str(self.file_path), data_only=True)
        sheets_data = {}
        
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                raise ValueError(f"工作表不存在: {sheet_name}。可用工作表: {wb.sheetnames}")
            sheets_to_read = [sheet_name]
        elif all_sheets:
            sheets_to_read = wb.sheetnames
        else:
            sheets_to_read = [wb.active.title]
        
        for name in sheets_to_read:
            ws = wb[name]
            rows = []
            for row in ws.iter_rows():
                row_data = []
                for cell in row:
                    value = cell.value
                    if value is None:
                        value = ""
                    elif isinstance(value, (int, float)):
                        value = str(value)
                    row_data.append(str(value))
                rows.append(row_data)
            sheets_data[name] = rows
        
        return sheets_data
    
    def _read_xls(self, sheet_name: Optional[str], all_sheets: bool) -> Dict[str, List[List[Any]]]:
        """读取 .xls 文件"""
        import xlrd
        
        wb = xlrd.open_workbook(str(self.file_path))
        sheets_data = {}
        
        if sheet_name:
            if sheet_name not in wb.sheet_names():
                raise ValueError(f"工作表不存在: {sheet_name}。可用工作表: {wb.sheet_names()}")
            sheets_to_read = [sheet_name]
        elif all_sheets:
            sheets_to_read = wb.sheet_names()
        else:
            sheets_to_read = [wb.sheet_by_index(0).name]
        
        for name in sheets_to_read:
            ws = wb.sheet_by_name(name)
            rows = []
            for row_idx in range(ws.nrows):
                row_data = [str(ws.cell_value(row_idx, col_idx)) for col_idx in range(ws.ncols)]
                rows.append(row_data)
            sheets_data[name] = rows
        
        return sheets_data
    
    def to_markdown(self, sheet_name: Optional[str] = None, all_sheets: bool = False) -> str:
        """将 Excel 文档转换为 Markdown 格式"""
        data = self.read(sheet_name, all_sheets)
        lines = []
        
        lines.append(f"# {data['file_name']}")
        lines.append("")
        
        for sheet_name, rows in data["sheets"].items():
            lines.append(f"## 工作表: {sheet_name}")
            lines.append("")
            
            if not rows:
                lines.append("(空表)")
                lines.append("")
                continue
            
            # 找到最大列数
            max_cols = max(len(row) for row in rows) if rows else 0
            
            if max_cols > 0:
                # 表头
                header = rows[0] if rows else []
                header = header + [""] * (max_cols - len(header))
                lines.append("| " + " | ".join(header) + " |")
                lines.append("| " + " | ".join(["---"] * max_cols) + " |")
                
                # 数据行
                for row in rows[1:]:
                    padded_row = row + [""] * (max_cols - len(row))
                    lines.append("| " + " | ".join(padded_row) + " |")
            
            lines.append("")
        
        return "\n".join(lines)
    
    def to_text(self, sheet_name: Optional[str] = None, all_sheets: bool = False) -> str:
        """将 Excel 文档转换为纯文本"""
        data = self.read(sheet_name, all_sheets)
        lines = []
        
        lines.append(f"文件: {data['file_name']}")
        lines.append("=" * 50)
        lines.append("")
        
        for sheet_name, rows in data["sheets"].items():
            lines.append(f"工作表: {sheet_name}")
            lines.append("-" * 30)
            
            for row in rows:
                lines.append(" | ".join(row))
            
            lines.append("")
        
        return "\n".join(lines)


def read_document(
    file_path: str,
    output_format: str = "markdown",
    sheet_name: Optional[str] = None,
    all_sheets: bool = False
) -> str:
    """
    读取文档并返回指定格式的内容
    
    Args:
        file_path: 文件路径
        output_format: 输出格式 (text, json, markdown)
        sheet_name: Excel 工作表名称
        all_sheets: 是否读取所有工作表
    
    Returns:
        格式化后的文档内容
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == '.docx':
        reader = WordReader(file_path)
        if output_format == "json":
            return json.dumps(reader.read(), ensure_ascii=False, indent=2)
        elif output_format == "text":
            return reader.to_text()
        else:
            return reader.to_markdown()
    
    elif suffix in ['.xlsx', '.xls']:
        reader = ExcelReader(file_path)
        if output_format == "json":
            return json.dumps(reader.read(sheet_name, all_sheets), ensure_ascii=False, indent=2)
        elif output_format == "text":
            return reader.to_text(sheet_name, all_sheets)
        else:
            return reader.to_markdown(sheet_name, all_sheets)
    
    else:
        raise ValueError(f"不支持的文件格式: {suffix}。支持的格式: .docx, .xlsx, .xls")


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description="读取 Word 和 Excel 文档内容",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    # 读取 Word 文档
    python doc_reader.py document.docx
    
    # 读取 Excel 并输出为 JSON
    python doc_reader.py data.xlsx --format json
    
    # 读取 Excel 指定工作表
    python doc_reader.py data.xlsx --sheet "Sheet1"
    
    # 读取 Excel 所有工作表
    python doc_reader.py data.xlsx --all-sheets
    
    # 输出到文件
    python doc_reader.py document.docx --output output.md
        """
    )
    
    parser.add_argument("file", nargs='?', help="要读取的文件路径 (.docx, .xlsx, .xls)")
    parser.add_argument(
        "--format", "-f",
        choices=["text", "json", "markdown"],
        default="markdown",
        help="输出格式 (默认: markdown)"
    )
    parser.add_argument(
        "--sheet", "-s",
        help="指定 Excel 工作表名称"
    )
    parser.add_argument(
        "--all-sheets", "-a",
        action="store_true",
        help="读取所有工作表 (仅 Excel)"
    )
    parser.add_argument(
        "--output", "-o",
        help="输出到文件"
    )
    parser.add_argument(
        "--check-deps",
        action="store_true",
        help="检查依赖库安装状态"
    )
    
    args = parser.parse_args()
    
    # 检查依赖
    if args.check_deps:
        deps = check_dependencies()
        print("依赖检查结果:")
        for dep, installed in deps.items():
            status = "✓ 已安装" if installed else "✗ 未安装"
            print(f"  {dep}: {status}")
        
        missing = [dep for dep, installed in deps.items() if not installed]
        if missing:
            print(f"\n{install_hint(missing)}")
            sys.exit(1)
        return
    
    # 检查是否提供了文件参数
    if not args.file:
        parser.print_help()
        sys.exit(1)
    
    # 检查必要依赖
    deps = check_dependencies()
    suffix = Path(args.file).suffix.lower()
    
    if suffix == '.docx' and not deps['python-docx']:
        print("错误: 读取 .docx 文件需要安装 python-docx")
        print(install_hint(['python-docx']))
        sys.exit(1)
    
    if suffix == '.xlsx' and not deps['openpyxl']:
        print("错误: 读取 .xlsx 文件需要安装 openpyxl")
        print(install_hint(['openpyxl']))
        sys.exit(1)
    
    if suffix == '.xls' and not deps['xlrd']:
        print("错误: 读取 .xls 文件需要安装 xlrd")
        print(install_hint(['xlrd']))
        sys.exit(1)
    
    try:
        result = read_document(
            args.file,
            output_format=args.format,
            sheet_name=args.sheet,
            all_sheets=args.all_sheets
        )
        
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(result)
            print(f"已输出到: {args.output}")
        else:
            print(result)
    
    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
